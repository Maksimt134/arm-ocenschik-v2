import os, sys, webbrowser, uvicorn
from fastapi import FastAPI, HTTPException
from fastapi.staticfiles import StaticFiles

app = FastAPI()

OKN_DATABASE = {

    "77:01:0001001:1023": {
      "id": "okn-knop",
      "cadastral_number": "77:01:0001001:1023",
      "name": "Усадьба А.Л. Кнопа (Главный дом с палатами)",
      "address": "г. Москва, Колпачный пер., д. 5, стр. 2",
      "area": 1240.5,
      "year_built": 1890,
      "floors": 3,
      "walls_material": "Кирпич с облицовкой",
      "okn_category": "Федеральное значение",
      "is_okn": True,
      "wear_pct": 32,
      "history": "Построена в 1890 году архитектором Б.В. Фрейденбергом по заказу промышленника Андрея Кнопа. Усадьба представляет собой уникальный образец неоготики с сохранившимися готическими залами, витражами и парадной лестницей. Включает палаты XVII века в основании здания.",
      "coordinates": [55.757833, 37.641944],
      "building_outline": [
        [55.758066, 37.641722],
        [55.758110, 37.641955],
        [55.757911, 37.642099],
        [55.757885, 37.641988],
        [55.757755, 37.642050],
        [55.757715, 37.641880],
        [55.757890, 37.641790],
        [55.757910, 37.641850],
        [55.758066, 37.641722]
      ],
      "bti_pdf_url": "#download-bti",
      "restrictions": [
        {
          "id": "restr-1",
          "title": "Охранная зона ОКН (Режим Р-1)",
          "description": "Запрещено любое новое строительство, изменение параметров существующих объектов, прокладка наземных инженерных коммуникаций. Разрешена только реставрация и консервация.",
          "severity": "high",
          "law_base": "Федеральный закон № 73-ФЗ, ст. 34"
        },
        {
          "id": "restr-2",
          "title": "Зона регулирования застройки (Режим Р-2)",
          "description": "Ограничение высотности новых зданий в буферной зоне до 12 метров. Требование сохранения исторических линий застройки и визуальных коридоров видимости.",
          "severity": "medium",
          "law_base": "Постановление Правительства Москвы № 121-ПП"
        },
        {
          "id": "restr-3",
          "title": "Предмет охраны внутреннего убранства",
          "description": "Обязательство сохранения лепного декора, дубовых панелей, изразцовых каминов конца XIX века. Любые работы подлежат согласованию с Мосгорнаследием.",
          "severity": "high",
          "law_base": "Охранное обязательство № 44-ОБ/2018"
        }
      ],
      "kkh_params": {
        "historical_weight": 4.5,
        "architectural_rarity": 4.8,
        "public_awareness": 4.2,
        "constraint_points": 3.8
      },
      "photos": [
        "https://images.unsplash.com/photo-1548625149-fc4a29cf7092?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1564507592224-2c266a2e4e46?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1510697723223-9ee82c817df0?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1577717903315-1691ae25ab3f?auto=format&fit=crop&w=800&q=80"
      ]
    },
    "77:01:0003020:4105": {
      "id": "okn-morozov",
      "cadastral_number": "77:01:0003020:4105",
      "name": "Особняк З.Г. Морозовой",
      "address": "г. Москва, ул. Спиридоновка, д. 17, стр. 1",
      "area": 2150.0,
      "year_built": 1898,
      "floors": 2,
      "walls_material": "Камень, кирпич",
      "okn_category": "Федеральное значение",
      "is_okn": True,
      "wear_pct": 25,
      "history": "Шедевр раннего московского модерна, построенный архитектором Ф.О. Шехтелем. В оформлении интерьеров принимал участие М.А. Врубель. В настоящее время является Домом приемов МИД РФ.",
      "coordinates": [55.760128, 37.595604],
      "building_outline": [
        [55.760311, 37.595288],
        [55.760380, 37.595500],
        [55.760200, 37.595600],
        [55.760250, 37.595850],
        [55.759988, 37.595911],
        [55.759920, 37.595700],
        [55.760050, 37.595600],
        [55.759888, 37.595499],
        [55.760311, 37.595288]
      ],
      "bti_pdf_url": "#download-bti",
      "restrictions": [
        {
          "id": "restr-morozov-1",
          "title": "Зона строгого охранного режима (Режим Р-1)",
          "description": "Полный запрет хозяйственной деятельности, не связанной с сохранением памятника.",
          "severity": "high",
          "law_base": "ФЗ № 73-ФЗ"
        }
      ],
      "kkh_params": {
        "historical_weight": 5.0,
        "architectural_rarity": 5.0,
        "public_awareness": 4.9,
        "constraint_points": 4.5
      },
      "photos": [
        "https://images.unsplash.com/photo-1580587771525-78b9dba3b469?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1600596542815-ffad4c1539a9?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1505843513577-22bb7d21e455?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1536647494541-69efdc7a5ec4?auto=format&fit=crop&w=800&q=80"
      ]
    },
    "77:01:0001037:1019": {
      "id": "okn-rossiya",
      "cadastral_number": "77:01:0001037:1019",
      "name": "Доходный дом страхового общества «Россия»",
      "address": "г. Москва, Сретенский бульвар, д. 6/1, стр. 1",
      "area": 24308.3,
      "year_built": 1901,
      "floors": 5,
      "walls_material": "Кирпичные",
      "okn_category": "Федеральное значение",
      "is_okn": True,
      "wear_pct": 30,
      "history": "Комплекс доходных домов, построенный по проекту Н.М. Проскурнина в 1899-1902 годах. Яркий памятник архитектуры в духе позднего итальянского ренессанса. Уникальный пример элитного жилья начала XX века с передовым для того времени техническим оснащением: собственные артезианские скважины, электростанция и система вентиляции.",
      "coordinates": [55.766333, 37.6335],
      "building_outline": [
        [55.766601, 37.632711],
        [55.766812, 37.633789],
        [55.766155, 37.634123],
        [55.765911, 37.633099],
        [55.766205, 37.632944],
        [55.766299, 37.633199],
        [55.766388, 37.632855],
        [55.766601, 37.632711]
      ],
      "bti_pdf_url": "/files/bti_rossiya_act.pdf",
      "restrictions": [
        {
          "id": "restr-rossiya-1",
          "title": "Охранная зона ОКН (Режим Р-1)",
          "description": "Запрещено любое новое строительство.",
          "severity": "high",
          "law_base": "ФЗ № 73-ФЗ"
        },
        {
          "id": "restr-rossiya-2",
          "title": "Зона регулирования застройки (Режим Р-2)",
          "description": "Ограничение высотности новых зданий в буферной зоне.",
          "severity": "medium",
          "law_base": "Постановление Правительства Москвы № 121-ПП"
        }
      ],
      "kkh_params": {
        "historical_weight": 5.0,
        "architectural_rarity": 4.9,
        "public_awareness": 5.0,
        "constraint_points": 4.2
      },
      "photos": [
        "https://images.unsplash.com/photo-1486406146926-c627a92ad1ab?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1523413651479-59cb1f1f6f9e?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1516156008625-3a9d6067fab5?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1503614472-8c93d56e92ce?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1497366216548-37526070297c?auto=format&fit=crop&w=800&q=80"
      ]
    },
    "77:01:0001045:201": {
      "id": "okn-ostozhenka",
      "cadastral_number": "77:01:0001045:201",
      "name": "Исторический особняк на Остоженке",
      "address": "ул. Остоженка, 53/2",
      "area": 19800.0,
      "year_built": 1903,
      "floors": 4,
      "walls_material": "Кирпич",
      "okn_category": "Федеральное значение",
      "is_okn": True,
      "wear_pct": 28,
      "history": "Историческое здание начала XX века, представляющее собой уникальный образец московского ампира с элементами раннего модерна. Фасад богато украшен лепниной, а во внутренних помещениях сохранились подлинные камины и дубовый паркет. В прошлом здесь располагались доходные квартиры для интеллигенции.",
      "coordinates": [55.7366, 37.5947],
      "building_outline": [
        [55.736822, 37.594255],
        [55.736933, 37.594788],
        [55.736412, 37.595088],
        [55.736311, 37.594555],
        [55.736822, 37.594255]
      ],
      "bti_pdf_url": "#",
      "restrictions": [
        {
          "id": "restr-ostozh-1",
          "title": "Охранная зона ОКН (Режим Р-1)",
          "description": "Запрет на изменение объемно-пространственных характеристик здания. Запрещена установка кондиционеров на главном фасаде.",
          "severity": "high",
          "law_base": "ФЗ № 73-ФЗ"
        }
      ],
      "kkh_params": {
        "historical_weight": 4.0,
        "architectural_rarity": 4.0,
        "public_awareness": 4.0,
        "constraint_points": 3.5
      },
      "photos": [
        "https://images.unsplash.com/photo-1519999482648-25049ddd37b1?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1524813686514-a57563d77965?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1502672260266-1c1c8f1ebcda?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1481277542470-605612bd2d61?auto=format&fit=crop&w=800&q=80"
      ]
    },
    "77:01:0001061:88": {
      "id": "okn-prechistenka",
      "cadastral_number": "77:01:0001061:88",
      "name": "Городская усадьба XVIII-XIX веков",
      "address": "ул. Пречистенка, 13",
      "area": 21500.0,
      "year_built": 1910,
      "floors": 4,
      "walls_material": "Кирпич",
      "okn_category": "Федеральное значение",
      "is_okn": True,
      "wear_pct": 25,
      "history": "Классическая московская усадьба, перестроенная после пожара 1812 года. Ансамбль включает главный дом с мезонином и флигели. Особую ценность представляет чугунная ограда и сохранившиеся интерьеры парадных залов с потолочными росписями.",
      "coordinates": [55.7420, 37.5960],
      "building_outline": [
        [55.742211, 37.595555],
        [55.742333, 37.596188],
        [55.741812, 37.596399],
        [55.741699, 37.595755],
        [55.742211, 37.595555]
      ],
      "bti_pdf_url": "#",
      "restrictions": [
        {
          "id": "restr-prechist-1",
          "title": "Предмет охраны",
          "description": "Обязательство сохранения анфиладной планировки и элементов первоначального декора. Любые работы подлежат согласованию.",
          "severity": "medium",
          "law_base": "Охранное обязательство"
        }
      ],
      "kkh_params": {
        "historical_weight": 4.5,
        "architectural_rarity": 4.2,
        "public_awareness": 4.1,
        "constraint_points": 3.7
      },
      "photos": [
        "https://images.unsplash.com/photo-1584622650111-993a426fbf0a?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1605810230434-7631ac76ec81?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1585421522853-c6258d550e58?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1591825729272-8705a41dbcd5?auto=format&fit=crop&w=800&q=80"
      ]
    },
    "77:01:0001074:340": {
      "id": "okn-myasnitskaya",
      "cadastral_number": "77:01:0001074:340",
      "name": "Дом-усадьба на Мясницкой",
      "address": "ул. Мясницкая, 22",
      "area": 25100.0,
      "year_built": 1898,
      "floors": 4,
      "walls_material": "Кирпич",
      "okn_category": "Региональное значение",
      "is_okn": True,
      "wear_pct": 32,
      "history": "Выдающийся памятник архитектуры зрелого классицизма. Главный дом с выразительным портиком и симметричными флигелями формирует красную линию улицы. В здании частично сохранилась первоначальная планировка.",
      "coordinates": [55.7620, 37.6360],
      "building_outline": [
        [55.762211, 37.635555],
        [55.762333, 37.636188],
        [55.761812, 37.636399],
        [55.761699, 37.635755],
        [55.762211, 37.635555]
      ],
      "bti_pdf_url": "#",
      "restrictions": [
        {
          "id": "restr-myasn-1",
          "title": "Зона строгого охранного режима",
          "description": "Полный запрет земляных работ в историческом контуре двора. Ограничения на ремонт инженерных систем.",
          "severity": "high",
          "law_base": "ФЗ № 73-ФЗ"
        }
      ],
      "kkh_params": {
        "historical_weight": 4.1,
        "architectural_rarity": 4.3,
        "public_awareness": 4.5,
        "constraint_points": 4.0
      },
      "photos": [
        "https://images.unsplash.com/photo-1572120360610-d971b9d7767c?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1568605114967-8130f3a36994?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1554995207-c18c203602cb?auto=format&fit=crop&w=800&q=80",
        "https://images.unsplash.com/photo-1590059379685-64f433ce10bd?auto=format&fit=crop&w=800&q=80"
      ]
    },
  }

ANALOGUES_DATABASE = [
    {
        "id": "analog-1", "address": "г. Москва, Хохловский пер., д. 7-9, стр. 2", "area": 1100.0, "year_built": 1905, "floors": 3, 
        "walls_material": "Кирпич", "is_okn": 1.0, "wear_pct": 35.0, "dist_metro_min": 10.0, "infrastructure_rate": 4.0, 
        "noise_rate": 2.0, "parking": 1.0, "view_rate": 4.0, "base_price": 540000000.0, "cyan_url": "", 
        "coordinates": [55.755915, 37.643668]
    },
    {
        "id": "analog-2", "address": "г. Москва, Подкопаевский пер., д. 4, стр. 1", "area": 1450.0, "year_built": 1880, "floors": 2, 
        "walls_material": "Кирпич с каменным цоколем", "is_okn": 1.0, "wear_pct": 40.0, "dist_metro_min": 12.0, "infrastructure_rate": 4.5, 
        "noise_rate": 1.0, "parking": 0.0, "view_rate": 3.5, "base_price": 680000000.0, "cyan_url": "", 
        "coordinates": [55.753737, 37.641261]
    },
    {
        "id": "analog-3", "address": "г. Москва, ул. Солянка, д. 12, стр. 3", "area": 1300.0, "year_built": 1912, "floors": 4, 
        "walls_material": "Кирпич железобетон", "is_okn": 0.0, "wear_pct": 28.0, "dist_metro_min": 5.0, "infrastructure_rate": 5.0, 
        "noise_rate": 4.0, "parking": 0.0, "view_rate": 4.0, "base_price": 710000000.0, "cyan_url": "", 
        "coordinates": [55.751254, 37.641117]
    },
    {
        "id": "analog-4", "address": "г. Москва, Покровский бульвар, д. 8, стр. 1", "area": 950.0, "year_built": 1895, "floors": 3, 
        "walls_material": "Кирпич", "is_okn": 1.0, "wear_pct": 20.0, "dist_metro_min": 14.0, "infrastructure_rate": 4.5, 
        "noise_rate": 3.0, "parking": 1.0, "view_rate": 5.0, "base_price": 620000000.0, "cyan_url": "", 
        "coordinates": [55.757328, 37.646561]
    },
    {
        "id": "analog-5", "address": "г. Москва, ул. Забелина, д. 3", "area": 1600.0, "year_built": 1900, "floors": 3, 
        "walls_material": "Кирпич", "is_okn": 1.0, "wear_pct": 45.0, "dist_metro_min": 2.0, "infrastructure_rate": 5.0, 
        "noise_rate": 4.5, "parking": 0.0, "view_rate": 4.5, "base_price": 750000000.0, "cyan_url": "", 
        "coordinates": [55.7551, 37.6391]
    },
    {
        "id": "analog-6", "address": "г. Москва, Старосадский пер., д. 9", "area": 1250.0, "year_built": 1908, "floors": 4, 
        "walls_material": "Кирпич", "is_okn": 1.0, "wear_pct": 30.0, "dist_metro_min": 5.0, "infrastructure_rate": 4.8, 
        "noise_rate": 2.0, "parking": 1.0, "view_rate": 4.0, "base_price": 590000000.0, "cyan_url": "", 
        "coordinates": [55.7565, 37.6398]
    },
    {
        "id": "analog-7", "address": "г. Москва, Малый Златоустинский пер., д. 4", "area": 1050.0, "year_built": 1890, "floors": 3, 
        "walls_material": "Кирпич", "is_okn": 1.0, "wear_pct": 33.0, "dist_metro_min": 6.0, "infrastructure_rate": 4.5, 
        "noise_rate": 2.5, "parking": 0.0, "view_rate": 3.5, "base_price": 490000000.0, "cyan_url": "", 
        "coordinates": [55.7588, 37.6325]
    },
    {
        "id": "analog-8", "address": "г. Москва, Колпачный пер., д. 10", "area": 1400.0, "year_built": 1902, "floors": 5, 
        "walls_material": "Кирпич", "is_okn": 0.0, "wear_pct": 28.0, "dist_metro_min": 8.0, "infrastructure_rate": 4.4, 
        "noise_rate": 3.0, "parking": 1.0, "view_rate": 4.0, "base_price": 630000000.0, "cyan_url": "", 
        "coordinates": [55.7595, 37.6445]
    },
    {
        "id": "analog-9", "address": "г. Москва, Чистопрудный бульвар, д. 12", "area": 1750.0, "year_built": 1910, "floors": 6, 
        "walls_material": "Смешанные", "is_okn": 1.0, "wear_pct": 31.0, "dist_metro_min": 4.0, "infrastructure_rate": 5.0, 
        "noise_rate": 3.5, "parking": 1.0, "view_rate": 4.5, "base_price": 790000000.0, "cyan_url": "", 
        "coordinates": [55.7608, 37.6435]
    },
    {
        "id": "analog-10", "address": "г. Москва, Покровский бульвар, д. 10", "area": 1150.0, "year_built": 1898, "floors": 3, 
        "walls_material": "Кирпич", "is_okn": 0.0, "wear_pct": 29.0, "dist_metro_min": 9.0, "infrastructure_rate": 4.2, 
        "noise_rate": 2.8, "parking": 0.0, "view_rate": 3.8, "base_price": 510000000.0, "cyan_url": "", 
        "coordinates": [55.7595, 37.6495]
    }
]

ROSSIYA_ANALOGUES_DATABASE = [
    {
        "id": "analogue-rossiya-1", "address": "г. Москва, Большой Харитоньевский пер., 10", "area": 12400.0, "year_built": 1900, "floors": 5, 
        "walls_material": "Кирпичные", "is_okn": 1.0, "wear_pct": 26.0, "dist_metro_min": 7.0, "infrastructure_rate": 4.5, 
        "noise_rate": 2.0, "parking": 1.0, "view_rate": 4.0, "base_price": 620000000.0, "cyan_url": "", 
        "coordinates": [55.7628, 37.6441]
    },
    {
        "id": "analogue-rossiya-2", "address": "г. Москва, Курсовой пер., 1", "area": 9800.0, "year_built": 1907, "floors": 6, 
        "walls_material": "Кирпичные с облицовкой", "is_okn": 1.0, "wear_pct": 22.0, "dist_metro_min": 5.0, "infrastructure_rate": 5.0, 
        "noise_rate": 1.5, "parking": 0.0, "view_rate": 4.5, "base_price": 710000000.0, "cyan_url": "", 
        "coordinates": [55.7425, 37.6048]
    },
    {
        "id": "analogue-rossiya-3", "address": "г. Москва, Пречистенка, 28", "area": 15200.0, "year_built": 1898, "floors": 5, 
        "walls_material": "Кирпичные", "is_okn": 1.0, "wear_pct": 31.0, "dist_metro_min": 8.0, "infrastructure_rate": 4.8, 
        "noise_rate": 2.5, "parking": 1.0, "view_rate": 4.0, "base_price": 580000000.0, "cyan_url": "", 
        "coordinates": [55.7410, 37.5930]
    },
    {
        "id": "analogue-rossiya-4", "address": "г. Москва, Арбат, 35", "area": 8700.0, "year_built": 1903, "floors": 6, 
        "walls_material": "Кирпичные", "is_okn": 1.0, "wear_pct": 28.0, "dist_metro_min": 4.0, "infrastructure_rate": 5.0, 
        "noise_rate": 3.5, "parking": 0.0, "view_rate": 3.8, "base_price": 495000000.0, "cyan_url": "", 
        "coordinates": [55.7495, 37.5910]
    },
    {
        "id": "analogue-rossiya-5", "address": "г. Москва, Гоголевский бульвар, 21", "area": 11300.0, "year_built": 1905, "floors": 5, 
        "walls_material": "Кирпичные с лепниной", "is_okn": 0.0, "wear_pct": 33.0, "dist_metro_min": 6.0, "infrastructure_rate": 4.7, 
        "noise_rate": 3.0, "parking": 1.0, "view_rate": 4.2, "base_price": 530000000.0, "cyan_url": "", 
        "coordinates": [55.7485, 37.6015]
    },
    {
        "id": "analogue-rossiya-6", "address": "г. Москва, Романов пер., 3", "area": 16800.0, "year_built": 1898, "floors": 6, 
        "walls_material": "Кирпичные", "is_okn": 1.0, "wear_pct": 29.0, "dist_metro_min": 5.0, "infrastructure_rate": 4.9, 
        "noise_rate": 2.0, "parking": 1.0, "view_rate": 4.5, "base_price": 650000000.0, "cyan_url": "", 
        "coordinates": [55.7538, 37.6095]
    },
    {
        "id": "analogue-rossiya-7", "address": "г. Москва, Мясницкая ул., 15", "area": 9200.0, "year_built": 1895, "floors": 5, 
        "walls_material": "Кирпичные", "is_okn": 0.0, "wear_pct": 35.0, "dist_metro_min": 6.0, "infrastructure_rate": 4.5, 
        "noise_rate": 3.0, "parking": 0.0, "view_rate": 3.5, "base_price": 460000000.0, "cyan_url": "", 
        "coordinates": [55.7618, 37.6335]
    },
    {
        "id": "analogue-rossiya-8", "address": "г. Москва, Покровка, 2/1", "area": 14100.0, "year_built": 1902, "floors": 6, 
        "walls_material": "Кирпичные", "is_okn": 1.0, "wear_pct": 24.0, "dist_metro_min": 8.0, "infrastructure_rate": 4.4, 
        "noise_rate": 2.8, "parking": 0.0, "view_rate": 3.8, "base_price": 590000000.0, "cyan_url": "", 
        "coordinates": [55.758784, 37.640108]
    },
    {
        "id": "analogue-rossiya-9", "address": "г. Москва, Тверская ул., 25", "area": 19800.0, "year_built": 1898, "floors": 5, 
        "walls_material": "Кирпичные с облицовкой", "is_okn": 0.0, "wear_pct": 32.0, "dist_metro_min": 3.0, "infrastructure_rate": 5.0, 
        "noise_rate": 4.0, "parking": 1.0, "view_rate": 5.0, "base_price": 412000000.0, "cyan_url": "", 
        "coordinates": [55.767425, 37.599723]
    },
    {
        "id": "analogue-rossiya-10", "address": "г. Москва, Б. Лубянка, 14", "area": 26700.0, "year_built": 1907, "floors": 6, 
        "walls_material": "Кирпичные", "is_okn": 1.0, "wear_pct": 27.0, "dist_metro_min": 6.0, "infrastructure_rate": 5.0, 
        "noise_rate": 3.0, "parking": 0.0, "view_rate": 4.0, "base_price": 529000000.0, "cyan_url": "", 
        "coordinates": [55.761803, 37.629471]
    },
    {
        "id": "analogue-rossiya-11", "address": "г. Москва, Новокузнецкая ул., 34", "area": 18400.0, "year_built": 1910, "floors": 5, 
        "walls_material": "Смешанные", "is_okn": 0.0, "wear_pct": 30.0, "dist_metro_min": 8.0, "infrastructure_rate": 4.0, 
        "noise_rate": 2.0, "parking": 1.0, "view_rate": 3.0, "base_price": 398000000.0, "cyan_url": "", 
        "coordinates": [55.732731, 37.629606]
    },
    {
        "id": "analogue-rossiya-12", "address": "г. Москва, Столешников пер., 11", "area": 23100.0, "year_built": 1902, "floors": 5, 
        "walls_material": "Кирпичные", "is_okn": 1.0, "wear_pct": 25.0, "dist_metro_min": 4.0, "infrastructure_rate": 5.0, 
        "noise_rate": 2.0, "parking": 0.0, "view_rate": 4.0, "base_price": 467000000.0, "cyan_url": "", 
        "coordinates": [55.761833, 37.614833]
    },
    {
        "id": "analogue-rossiya-13", "address": "г. Москва, Мясницкая ул., 22", "area": 25400.0, "year_built": 1896, "floors": 6, 
        "walls_material": "Кирпичные с лепниной", "is_okn": 0.0, "wear_pct": 34.0, "dist_metro_min": 5.0, "infrastructure_rate": 5.0, 
        "noise_rate": 4.0, "parking": 1.0, "view_rate": 4.0, "base_price": 445000000.0, "cyan_url": "", 
        "coordinates": [55.762145, 37.634470]
    },
    {
        "id": "analogue-rossiya-14", "address": "г. Москва, Подкопаевский пер., 4, стр. 1", "area": 14500.0, "year_built": 1880, "floors": 5, 
        "walls_material": "Кирпичные", "is_okn": 1.0, "wear_pct": 29.0, "dist_metro_min": 12.0, "infrastructure_rate": 4.5, 
        "noise_rate": 1.0, "parking": 0.0, "view_rate": 3.5, "base_price": 680000000.0, "cyan_url": "", 
        "coordinates": [55.753737, 37.641261]
    }
]

@app.get("/api/search")
def search(q: str, type: str = "cadastral"):
    for cad, data in OKN_DATABASE.items():
        if q.lower() in cad.lower() or q.lower() in data["address"].lower():
            return {"status": "found", "cadastral": cad, "id": data["id"]}
    raise HTTPException(status_code=404, detail="Не найдено")

@app.get("/api/passport/{obj_id}")
def get_passport(obj_id: str):
    for cad, data in OKN_DATABASE.items():
        if data["id"] == obj_id or cad == obj_id:
            return data
    raise HTTPException(status_code=404, detail="Не найдено")

@app.get("/api/analogues")
def get_analogues(target_id: str, okn_filter: str = "all", area_tolerance: float = 0.3):
    if target_id == "okn-rossiya":
        return ROSSIYA_ANALOGUES_DATABASE
    return ANALOGUES_DATABASE

@app.post("/api/generate-report")
def generate_report(payload: dict):
    from fastapi.responses import Response
    try:
        from docx import Document
        from io import BytesIO
        
        okn = payload.get("okn", {})
        results = payload.get("results", {})
        
        okn_name = okn.get("name", "Неизвестный объект")
        cadastral = okn.get("cadastral_number", "—")
        area = okn.get("area", 0)
        
        total_val = results.get("total_value", 0)
        w_comp = results.get("comp_weight", 0) * 100
        w_inc = results.get("income_weight", 0) * 100
        w_cost = results.get("cost_weight", 0) * 100
        min_conf = results.get("min_confidence", 0)
        max_conf = results.get("max_confidence", 0)
        
        document = Document()
        document.add_heading('Отчет об оценке объекта недвижимости', 0)
        
        document.add_heading('Общие сведения об объекте', level=1)
        document.add_paragraph(f"Объект оценки: {okn_name}")
        document.add_paragraph(f"Кадастровый номер: {cadastral}")
        document.add_paragraph(f"Площадь: {area} кв.м")
        
        document.add_heading('Итоговые результаты', level=1)
        p_total = document.add_paragraph("Итоговая рыночная стоимость: ")
        p_total.add_run(f"{total_val:,.2f} руб.").bold = True
        
        document.add_paragraph("Доверительный интервал (±5%):")
        document.add_paragraph(f"Минимальная граница: {min_conf:,.2f} руб.", style='List Bullet')
        document.add_paragraph(f"Максимальная граница: {max_conf:,.2f} руб.", style='List Bullet')
        
        document.add_heading('Весовое согласование подходов', level=1)
        document.add_paragraph(f"Сравнительный подход: {w_comp:.0f}%", style='List Bullet')
        document.add_paragraph(f"Доходный подход: {w_inc:.0f}%", style='List Bullet')
        document.add_paragraph(f"Затратный подход: {w_cost:.0f}%", style='List Bullet')
        
        document.add_paragraph("\nОтчет сгенерирован автоматически системой АРМ Оценщика.")
        
        f = BytesIO()
        document.save(f)
        f.seek(0)
        content = f.read()
    except Exception as e:
        content = f"Отчет об оценке\nОбъект: {payload.get('okn', {}).get('name', 'Неизвестно')}\nОшибка генерации: {str(e)}".encode('utf-8')
        
    return Response(
        content=content, 
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

def get_resource_path():
    if getattr(sys, 'frozen', False):
        return os.path.join(sys._MEIPASS, 'frontend', 'dist')
    return os.path.join(os.path.dirname(__file__), '..', 'frontend', 'dist')

static_dir = get_resource_path()
if os.path.exists(static_dir):
    app.mount("/", StaticFiles(directory=static_dir, html=True), name="static")

if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8000)
